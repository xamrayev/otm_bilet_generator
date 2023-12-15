from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import random


def biletlar_ru(bilet_soni=int, savollar=list, fan=str,semestr=str, kafedra=str, tuzuvchi=str, zav_kaf=str):
    '''
    bilet_soni => int, savollar => list, fan=>str, semestr=>int, kafedra=>str, tuzuvchi=>str, zav_kaf=>str,
    tayyor savollarni random metodida bilet_soni marta biletlarga ajratib beradi.
    '''
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    schetchik=0
    for i in range(bilet_soni):
        p_v1 = document.add_paragraph()
        p_v2 = document.add_paragraph()
        p_v3 = document.add_paragraph()
        v1 = p_v1.add_run('\nНаманганский инженерно-строительный институт\n')
        v2 = p_v1.add_run(f'Кафедра «{kafedra}» Билеты для проведения промежуточной работы\n')
        v3 = p_v1.add_run(f'по дисциплине «{fan}» ({semestr}-семестр)\n')
        v4 = p_v1.add_run(f'ВАРИАНТ № {i+1}')
        v5 = p_v2.add_run(f"1) {savollar[schetchik+0]}\n")
        v6 = p_v2.add_run(f"2) {savollar[schetchik+1]}\n")
        v7 = p_v2.add_run(f"3) {savollar[schetchik+2]}\n")
        v8 = p_v2.add_run(f"4) {savollar[schetchik+3]}\n")
        v9 = p_v2.add_run(f"5) {savollar[schetchik+4]}")
        schetchik+=5
        # v5 = p_v2.add_run(f"1) {random.choice(savollar)}\n")
        # v6 = p_v2.add_run(f"2) {random.choice(savollar)}\n")
        # v7 = p_v2.add_run(f"3) {random.choice(savollar)}\n")
        # v8 = p_v2.add_run(f"4) {random.choice(savollar)}\n")
        # v9 = p_v2.add_run(f"5) {random.choice(savollar)}")
        p_v1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_v2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_v3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        v1.bold = True
        v2.bold = True
        v3.bold = True
        v4.bold = True
        table = document.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cell1 = table.rows[0].cells
        hdr_cell1[0].text = 'Составитель:'
        hdr_cell1[1].text = ''
        hdr_cell1[2].text = f'{tuzuvchi}'
        hdr_cell2 = table.rows[1].cells
        hdr_cell2[0].text = 'Заведующий кафедрой:'
        hdr_cell2[1].text = ''
        hdr_cell2[2].text = f'{zav_kaf}'
        
    filename = fan+"_biletlar1.docx"
    document.save(filename)
    print("100.0 %")
