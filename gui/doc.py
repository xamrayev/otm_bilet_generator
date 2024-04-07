from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def to_doc(
        til=str, 
        testlar = dict, 
        savol_soni = int, 
        variant=int, 
        kafedra_nomi =str,
        guruh = str,
        fan = str,
        semestr = int,
        tuzuvchi = str,
        zav_kaf = str,
        ):
    new_directory = f"{fan}"
    current_directory = os.getcwd()
    new_directory_path = os.path.join(current_directory, new_directory)
    if not os.path.exists(new_directory_path):
        os.makedirs(new_directory_path)

    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    p_v1 = document.add_paragraph()
    p_v2 = document.add_paragraph()
    p_v3 = document.add_paragraph()
    p_v4 = document.add_paragraph()
    if til=="ru":
        v1 = p_v1.add_run('НАМАНГАНСКИЙ ИНЖЕНЕРНО-СТРОИТЕЛЬНЫЙ ИНСТИТУТ\n')
        p_v1.add_run("\n")
        v2 = p_v1.add_run(f'КАФЕДРА «{kafedra_nomi}»\n \n \n БИЛЕТЫ ДЛЯ ПРОВЕДЕНИЯ ПРОМЕЖУТОЧНОЙ РАБОТЫ\n')
        p_v1.add_run("\n"*3)
        v3 = p_v1.add_run(f'ПО ПРЕДМЕТУ «{fan}»\n({semestr}-семестр)\n')
        p_v1.add_run("\n")
        v4 = p_v1.add_run(f'СТУДЕНТКА(КИ) ГР. {guruh}\n')
        p_v1.add_run("\n"*2)
        v5 = p_v1.add_run('_______________________________________________')
        p_v1.add_run("\n")
        v6 = p_v1.add_run(f'Вариант - {variant}')
        p_v1.add_run("\n")
    else:
        v1 = p_v1.add_run('Namangan muhandislik - qurilish instituti\n')
        p_v1.add_run("\n")
        v2 = p_v1.add_run(f'«{kafedra_nomi}» kafedrasi\n \n \n Oraliq nazorat uchun test\n')
        p_v1.add_run("\n"*3)
        v3 = p_v1.add_run(f'«{fan}» fanidan\n({semestr}-semestr uchun)\n')
        p_v1.add_run("\n")
        v4 = p_v1.add_run(f'{guruh} - guruh talabasi\n')
        p_v1.add_run("\n"*2)
        v5 = p_v1.add_run('_______________________________________________')
        p_v1.add_run("\n"*2)
        v6 = p_v1.add_run(f'Variant - {variant}')
        p_v1.add_run("\n")

    p_v1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_v2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_v3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    v1.bold = True
    v2.bold = True
    v3.bold = True
    v4.bold = True
    v5.bold = True

    table = document.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cell1 = table.rows[0].cells

    if til=="ru":
        hdr_cell1[0].text = 'Составитель:'
        hdr_cell1[1].text = ''
        hdr_cell1[2].text = f'{tuzuvchi}'
        hdr_cell2 = table.rows[1].cells
        hdr_cell2[0].text = 'Заведующий кафедрой:'
        hdr_cell2[1].text = ''
        hdr_cell2[2].text = f'{zav_kaf}'
        document.add_paragraph("\n"*2)
        nab_ball = document.add_paragraph("Набранный балл: _______________________")
    else:
        hdr_cell1[0].text = 'Tuzuvchi:'
        hdr_cell1[1].text = ''
        hdr_cell1[2].text = f'{tuzuvchi}'
        hdr_cell2 = table.rows[1].cells
        hdr_cell2[0].text = 'Kafedra mudiri:'
        hdr_cell2[1].text = ''
        hdr_cell2[2].text = f'{zav_kaf}'
        document.add_paragraph("\n"*2)
        nab_ball = document.add_paragraph("To'plangan ball: _______________________")
    nab_ball.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()
    for savol in range(1,savol_soni+1):

        document.add_paragraph(f"{savol}) {testlar[savol]['savol']}")

        # tablitsa
        table2 = document.add_table(rows=2, cols=2)
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        savol_cell1 = table2.rows[0].cells
        savol_cell1[0].text = f"\ta){testlar[savol]['javob'][0]}"
        savol_cell1[1].text = f"\tb){testlar[savol]['javob'][1]}"
        savol_cell2 = table2.rows[1].cells
        savol_cell2[0].text = f"\tc){testlar[savol]['javob'][2]}"
        savol_cell2[1].text = f"\td){testlar[savol]['javob'][3]}"

        # ketma - ket
        # document.add_paragraph(f"\ta){testlar[savol]['javob'][0]}")
        # document.add_paragraph(f"\tb){testlar[savol]['javob'][1]}")
        # document.add_paragraph(f"\tc){testlar[savol]['javob'][2]}")
        # document.add_paragraph(f"\td){testlar[savol]['javob'][3]}\n")
    document.add_page_break()

    # javob sahifa
    
    savollar = savol_soni
    savol_q = savollar%10

    if savol_q !=0:
        savollar = (savollar//10+1)*10

    savol_row = int((savollar // 10*10)/5)
    savol_ort = savollar % 10


    table4 = document.add_table(rows=savol_row, cols=10, style="Table Grid")
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    first_row = table4.rows[0]

    for i, cell in enumerate(first_row.cells, start=1):
        cell.text = str(i)

    sum=10

    for i in range(2, len(table4.rows),2):
        row = table4.rows[i]
        for j, cell in enumerate(row.cells, start=1):
            sum+=1
            cell.text = str(sum) 

    filename = f"{fan[:5]}_"+f"variant_{variant}.docx"
    file_path = os.path.join(new_directory_path, filename)

    document.save(file_path)
    # print(f"{variant}-variant 100.0 %")

# variantlar = {1: {1: {'savol': 'Axsikent qayerda', 'javob': ['Namanganda', 'Sardobada', 'Chorsuda', 'Sheyda']}, 2: {'savol': 'Bugun qanaqa kun', 'javob': ['Bugun', 'Kecha', 'Ertalab', 'Ozonda']}, 3: {'savol': 'Poytaxtimiz', 'javob': ['Anjan', 'Baliqchi', "Qo'qan", 'Toshkent']}, 4: {'savol': 'Maymun', 'javob': ['kiyik', 'Orangutan', 'shoxli', 'quyon']}}, 2: {1: {'savol': 'Bugun qanaqa kun', 'javob': ['Ertalab', 'Bugun', 'Kecha', 'Ozonda']}, 2: {'savol': 'Poytaxtimiz', 'javob': ['Baliqchi', "Qo'qan", 'Toshkent', 'Anjan']}, 3: {'savol': 'Axsikent qayerda', 'javob': ['Namanganda', 'Sheyda', 'Chorsuda', 'Sardobada']}, 4: {'savol': 'Maymun', 'javob': ['shoxli', 'Orangutan', 'quyon', 'kiyik']}}, 3: {1: {'savol': 'Axsikent qayerda', 'javob': ['Sardobada', 'Namanganda', 'Sheyda', 'Chorsuda']}, 2: {'savol': 'Poytaxtimiz', 'javob': ['Baliqchi', 'Toshkent', 'Anjan', "Qo'qan"]}, 3: {'savol': 'Bugun qanaqa kun', 'javob': ['Bugun', 'Kecha', 'Ozonda', 'Ertalab']}, 4: {'savol': 'Maymun', 'javob': ['shoxli', 'Orangutan', 'quyon', 'kiyik']}}, 4: {1: {'savol': 'Maymun', 'javob': ['Orangutan', 'kiyik', 'shoxli', 'quyon']}, 2: {'savol': 'Axsikent qayerda', 'javob': ['Sheyda', 'Sardobada', 'Chorsuda', 'Namanganda']}, 3: {'savol': 'Poytaxtimiz', 'javob': ['Baliqchi', 'Toshkent', "Qo'qan", 'Anjan']}, 4: {'savol': 'Bugun qanaqa kun', 'javob': ['Ertalab', 'Kecha', 'Ozonda', 'Bugun']}}, 5: {1: {'savol': 'Maymun', 'javob': ['shoxli', 'Orangutan', 'kiyik', 'quyon']}, 2: {'savol': 'Axsikent qayerda', 'javob': ['Namanganda', 'Sardobada', 'Sheyda', 'Chorsuda']}, 3: {'savol': 'Poytaxtimiz', 'javob': ['Toshkent', 'Anjan', "Qo'qan", 'Baliqchi']}, 4: {'savol': 'Bugun qanaqa kun', 'javob': ['Ozonda', 'Ertalab', 'Kecha', 'Bugun']}}}

def test_to_doc(
        variantlar=dict,
        til = str,
        kafedra_nomi=str,
        guruh = str,
        fan = str,
        semestr = int,
        tuzuvchi = str,
        zav_kaf = str
        ):
    for i in range(1,len(variantlar.keys())+1):
        savollar = len(variantlar[i].keys())
        to_doc(til=til, 
                testlar=variantlar[i], 
                savol_soni=savollar, 
                variant=i,
                kafedra_nomi =kafedra_nomi,
                guruh = guruh,
                fan = fan,
                semestr = semestr,
                tuzuvchi = tuzuvchi,
                zav_kaf = zav_kaf
            )

    # for savol in range(1,savollar):
    #     print(variantlar[i][savol]['savol'])
    #     print(variantlar[i][savol]['javob'][0])
    #     print(variantlar[i][savol]['javob'][1])
    #     print(variantlar[i][savol]['javob'][2])
    #     print(variantlar[i][savol]['javob'][3])

    #     print("*"*10)
    # print("*"*20)

