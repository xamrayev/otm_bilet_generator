import openpyxl

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import random

def savollar_from_xls(filename=str):
    '''
    Exel faylidan ma'lumotlarni olish va ular asosida biletlar tuzish uchun ishlatiladi.
    funksiya file = file_to_path metodi da ishlaydi
    file_to_path o'rniga savollar joylashgan xls fayliga yo'l ko'rsating
    '''
    workbook = openpyxl.load_workbook(filename)
    sheet = workbook.active
    first_column_data = []
    for row in sheet.iter_rows():
        first_column_data.append(row[1].value)
    if sheet.max_row > 1:
        first_column_data = first_column_data[1:]
    return first_column_data

def biletlar_ru( guruh=str,bilet_soni=int, savollar=list, fan=str,semestr=str, kafedra=str, tuzuvchi=str, zav_kaf=str, savollar_soni=int):
    '''
    bilet_soni => int, savollar => list, fan=>str, semestr=>int, kafedra=>str, tuzuvchi=>str, zav_kaf=>str,
    tayyor savollarni random metodida bilet_soni marta biletlarga ajratib beradi.
    '''
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    for i in range(bilet_soni):
        p_v1 = document.add_paragraph()
        p_v2 = document.add_paragraph()
        p_v3 = document.add_paragraph()
        
        v1 = p_v1.add_run('Namangan muhandislik - qurilish instituti\n')
        v2 = p_v1.add_run(f'«{kafedra}» kafedrasi \n Oraliq nazorat uchun savollar\n')
        v2 = p_v1.add_run(f'«{fan}» fanidan {guruh} talabasi ({semestr}-semestr uchun)\n')
        v3 = p_v1.add_run(f'Oraliq nazorat savollari\n')
        v4 = p_v1.add_run(f'{i+1} - variant')
       
        for i in range(savollar_soni):
            p_v2.add_run(f"{i+1}) {random.choice(savollar)}\n")

        
        p_v1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_v2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_v3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        v1.bold = True
        v2.bold = True
        v3.bold = True
        v4.bold = True
        table = document.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cell1 = table.rows[0].cells
        
        hdr_cell1[0].text = 'Tuzuvchi:'
        hdr_cell1[1].text = ''
        hdr_cell1[2].text = f'{tuzuvchi}'
        hdr_cell2 = table.rows[1].cells
        hdr_cell2[0].text = 'Kafedra mudiri:'
        hdr_cell2[1].text = ''
        hdr_cell2[2].text = f'{zav_kaf}'
        # p_v3.add_run("\n")
    filename = fan+"_biletlar.docx"
    document.save(filename)
