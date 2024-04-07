from test_2_list import tayyorlash
from savol_2_doc import savollar_from_xls

from flask import Flask, render_template, send_file, request, jsonify, make_response

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


from io import BytesIO
from zipfile import ZipFile
import random


def generate_docx_files(
        kafedra_nomi=str,
        guruh = str,
        fan = str,
        semestr = int,
        tuzuvchi = str,
        zav_kaf = str,
        kerakli_variantlar_soni = int,
        variantda_test_soni = int,
        fayl = str,
        ):
   
    natija = tayyorlash(test_fayl=fayl, kerakli_variantlar_soni=kerakli_variantlar_soni,variantda_test_soni=variantda_test_soni)
    variantlar = natija[0]
    zip_buffer = BytesIO()
    with ZipFile(zip_buffer, 'w') as zip_file:
        for i in range(1,len(variantlar.keys())+1):
            variant=i
            savollar = len(variantlar[i].keys())
            savol_soni = len(variantlar[i].keys())
            testlar=variantlar[i]
            document = Document()
            style = document.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)
            p_v1 = document.add_paragraph()
            p_v2 = document.add_paragraph()
            p_v3 = document.add_paragraph()
            p_v4 = document.add_paragraph()
            
            v1 = p_v1.add_run('Namangan muhandislik - qurilish instituti\n')
            p_v1.add_run("\n")
            v2 = p_v1.add_run(f'«{kafedra_nomi}» kafedrasi\n \n \n Oraliq nazorat uchun test\n')
            p_v1.add_run("\n"*3)
            v3 = p_v1.add_run(f'«{fan}» fanidan\n({semestr}-semestr uchun)\n')
            p_v1.add_run("\n")
            v4 = p_v1.add_run(f'{guruh} - guruh talabasi\n')
            p_v1.add_run("\n"*2)
            v5 = p_v1.add_run('_______________________________________________')
            p_v1.add_run("\n"*2)
            v6 = p_v1.add_run(f'Variant - {variant}')
            p_v1.add_run("\n")

            p_v1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_v2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_v3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            v1.bold = True
            v2.bold = True
            v3.bold = True
            v4.bold = True
            v5.bold = True

            table = document.add_table(rows=2, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cell1 = table.rows[0].cells

        
            hdr_cell1[0].text = 'Tuzuvchi:'
            hdr_cell1[1].text = ''
            hdr_cell1[2].text = f'{tuzuvchi}'
            hdr_cell2 = table.rows[1].cells
            hdr_cell2[0].text = 'Kafedra mudiri:'
            hdr_cell2[1].text = ''
            hdr_cell2[2].text = f'{zav_kaf}'
            document.add_paragraph("\n"*2)
            nab_ball = document.add_paragraph("To'plangan ball: _______________________")
            nab_ball.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_page_break()

            for savol in range(1,savol_soni+1):

                document.add_paragraph(f"{savol}) {testlar[savol]['savol']}")

                # tablitsa
                table2 = document.add_table(rows=2, cols=2)
                table2.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                savol_cell1 = table2.rows[0].cells
                savol_cell1[0].text = f"\ta){testlar[savol]['javob'][0]}"
                savol_cell1[1].text = f"\tb){testlar[savol]['javob'][1]}"
                savol_cell2 = table2.rows[1].cells
                savol_cell2[0].text = f"\tc){testlar[savol]['javob'][2]}"
                savol_cell2[1].text = f"\td){testlar[savol]['javob'][3]}"


            document.add_page_break()

            # javob sahifa
            
            savollar = savol_soni
            savol_q = savollar%10

            if savol_q !=0:
                savollar = (savollar//10+1)*10

            savol_row = int((savollar // 10*10)/5)

            table4 = document.add_table(rows=savol_row, cols=10, style="Table Grid")
            table4.alignment = WD_TABLE_ALIGNMENT.CENTER
            first_row = table4.rows[0]

            for i, cell in enumerate(first_row.cells, start=1):
                cell.text = str(i)

            sum=10

            for i in range(2, len(table4.rows),2):
                row = table4.rows[i]
                for j, cell in enumerate(row.cells, start=1):
                    sum+=1
                    cell.text = str(sum) 
            
            filename = f"{fan[:5]}_"+f"variant_{variant}.docx"

            with BytesIO() as docx_buffer:
                        document.save(docx_buffer)
                        docx_buffer.seek(0)
                        zip_file.writestr(filename, docx_buffer.read())
        javoblar =  natija[1]
        document2 = Document()
        style = document2.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        j_v1 = document2.add_paragraph()

        j1 = j_v1.add_run('Variantlar uchun javoblar\n')
        j_v1.add_run("\n")
        j_v1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        j1.bold = True
        for v in range(1,len(javoblar.keys())+1):
            j_v1.add_run(f'{v} variant javoblari\n')
            for s in range(1,len(javoblar[v].keys())+1):
                javob = javoblar[v][s]['variant'].index(0)
                match javob:
                    case 0: 
                        j = "A"
                    case 1: 
                        j = "B"
                    case 2: 
                        j = "C"
                    case 3: 
                        j = "D"

                j_v1.add_run(f"{s} - {j}\n")
        

        filename2 = f"{fan}_javoblar.docx"

        with BytesIO() as docx_buffer:
            document2.save(docx_buffer)
            docx_buffer.seek(0)
            zip_file.writestr(filename2, docx_buffer.read())
    zip_buffer.seek(0)
    return zip_buffer.read()

def generate_docx_files_ru(
        kafedra_nomi=str,
        guruh = str,
        fan = str,
        semestr = int,
        tuzuvchi = str,
        zav_kaf = str,
        kerakli_variantlar_soni = int,
        variantda_test_soni = int,
        fayl = str,
        ):
    natija = tayyorlash(test_fayl=fayl, kerakli_variantlar_soni=kerakli_variantlar_soni,variantda_test_soni=variantda_test_soni)
    variantlar = natija[0]
    # print(type(len(variantlar.keys())))
    # print((len(variantlar.keys())))

    zip_buffer = BytesIO()
    
    with ZipFile(zip_buffer, 'w') as zip_file:
        
        for i in range(1,len(variantlar.keys())+1):
            variant=i
            savollar = len(variantlar[i].keys())
            savol_soni = len(variantlar[i].keys())
            testlar=variantlar[i]
            document = Document()
            style = document.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)
            p_v1 = document.add_paragraph()
            p_v2 = document.add_paragraph()
            p_v3 = document.add_paragraph()
            p_v4 = document.add_paragraph()
            
            v1 = p_v1.add_run('НАМАНГАНСКИЙ ИНЖЕНЕРНО-СТРОИТЕЛЬНЫЙ ИНСТИТУТ\n')
            p_v1.add_run("\n")
            v2 = p_v1.add_run(f'КАФЕДРА «{kafedra_nomi}»\n \n \n БИЛЕТЫ ДЛЯ ПРОВЕДЕНИЯ ПРОМЕЖУТОЧНОЙ РАБОТЫ\n')
            p_v1.add_run("\n"*3)
            v3 = p_v1.add_run(f'ПО ПРЕДМЕТУ «{fan}»\n({semestr}-семестр)\n')
            p_v1.add_run("\n")
            v4 = p_v1.add_run(f'СТУДЕНТКА(КИ) ГР. {guruh}\n')
            p_v1.add_run("\n"*2)
            v5 = p_v1.add_run('_______________________________________________')
            p_v1.add_run("\n")
            v6 = p_v1.add_run(f'Вариант - {variant}')
            p_v1.add_run("\n")

            p_v1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_v2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_v3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            v1.bold = True
            v2.bold = True
            v3.bold = True
            v4.bold = True
            v5.bold = True

            table = document.add_table(rows=2, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cell1 = table.rows[0].cells

        
            hdr_cell1[0].text = 'Составитель:'
            hdr_cell1[1].text = ''
            hdr_cell1[2].text = f'{tuzuvchi}'
            hdr_cell2 = table.rows[1].cells
            hdr_cell2[0].text = 'Заведующий кафедрой:'
            hdr_cell2[1].text = ''
            hdr_cell2[2].text = f'{zav_kaf}'
            document.add_paragraph("\n"*2)
            nab_ball = document.add_paragraph("Набранный балл: _______________________")
            nab_ball.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_page_break()

            for savol in range(1,savol_soni+1):
                document.add_paragraph(f"{savol}) {testlar[savol]['savol']}")

                # tablitsa
                table2 = document.add_table(rows=2, cols=2)
                table2.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                savol_cell1 = table2.rows[0].cells
                savol_cell1[0].text = f"\ta){testlar[savol]['javob'][0]}"
                savol_cell1[1].text = f"\tb){testlar[savol]['javob'][1]}"
                savol_cell2 = table2.rows[1].cells
                savol_cell2[0].text = f"\tc){testlar[savol]['javob'][2]}"
                savol_cell2[1].text = f"\td){testlar[savol]['javob'][3]}"


            document.add_page_break()

            # javob sahifa
            
            savollar = savol_soni
            savol_q = savollar%10

            if savol_q !=0:
                savollar = (savollar//10+1)*10

            savol_row = int((savollar // 10*10)/5)

            table4 = document.add_table(rows=savol_row, cols=10, style="Table Grid")
            table4.alignment = WD_TABLE_ALIGNMENT.CENTER
            first_row = table4.rows[0]

            for i, cell in enumerate(first_row.cells, start=1):
                # print(i,cell)
                cell.text = str(i)

            sum=10
            
            for i in range(2, len(table4.rows),2):
                row = table4.rows[i]
                for j, cell in enumerate(row.cells, start=1):
                    sum+=1
                    cell.text = str(sum) 
            
            filename = f"{fan[:5]}_"+f"вариант_{variant}.docx"

            with BytesIO() as docx_buffer:
                document.save(docx_buffer)
                docx_buffer.seek(0)
                zip_file.writestr(filename, docx_buffer.read())

        javoblar =  natija[1]
        document2 = Document()
        style = document2.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        j_v1 = document2.add_paragraph()

        j1 = j_v1.add_run('Ответы на вопросы\n')
        j_v1.add_run("\n")
        j_v1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        j1.bold = True
        for v in range(1,len(javoblar.keys())+1):
            j_v1.add_run(f'Ответы на вариант - {v}\n')
            for s in range(1,len(javoblar[v].keys())+1):
                javob = javoblar[v][s]['variant'].index(0)
                match javob:
                    case 0: 
                        j = "A"
                    case 1: 
                        j = "B"
                    case 2: 
                        j = "C"
                    case 3: 
                        j = "D"

                j_v1.add_run(f"{s} - {j}\n")
        

        filename2 = f"{fan}_ответы.docx"

        with BytesIO() as docx_buffer:
            document2.save(docx_buffer)
            docx_buffer.seek(0)
            zip_file.writestr(filename2, docx_buffer.read())
    zip_buffer.seek(0)
    return zip_buffer.read()

app = Flask(__name__, static_url_path='/static', static_folder='static')

@app.route('/')
def index():
    return render_template('index.html')  # Render the HTML form

@app.route('/ru')
def index_2():
    return render_template('index_2.html')  # Render the HTML form


@app.route('/generate_docx_ru', methods=['POST'])
def generate_docx_ru():
    try:
        # Get form data
        kafedra_nomi = request.form['kafedra_nomi']
        fan = request.form['fanimiz']
        semestr = request.form['semestr']
        tuzuvchi = request.form['tuzuvchi']
        kafedra_mudiri = request.form['kafedra_mudiri']
        bilet_soni = int(request.form['bilet_soni'])
        savollar_soni = int(request.form['savollar_soni'])
        # Get uploaded Excel file
        savol_file = request.files['savol_file']
        if savol_file.filename != '':
            savollar = savollar_from_xls(savol_file)
        
        document = Document()
        style = document.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)


        for i in range(bilet_soni):
            p_v1 = document.add_paragraph()
            p_v2 = document.add_paragraph()
            p_v3 = document.add_paragraph()
            
            v1 = p_v1.add_run('\nНаманганский инженерно-строительный институт\n')
            v2 = p_v1.add_run(f'Кафедра «{kafedra_nomi}» Билеты для проведения промежуточной работы\n')
            v3 = p_v1.add_run(f'по дисциплине «{fan}» ({semestr}-семестр)\n')
            v4 = p_v1.add_run(f'ВАРИАНТ № {i+1}')
        
            for i in range(savollar_soni):
                p_v2.add_run(f"{i+1}) {random.choice(savollar)}\n")

            
            p_v1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_v2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_v3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            v1.bold = True
            v2.bold = True
            v3.bold = True
            v4.bold = True
            table = document.add_table(rows=2, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cell1 = table.rows[0].cells
            
            hdr_cell1[0].text = 'Составитель:'
            hdr_cell1[1].text = ''
            hdr_cell1[2].text = f'{tuzuvchi}'
            hdr_cell2 = table.rows[1].cells
            hdr_cell2[0].text = 'Заведующий кафедрой:'
            hdr_cell2[1].text = ''
            hdr_cell2[2].text = f'{kafedra_mudiri}'
        
        
        output = BytesIO()
        document.save(output)
        output.seek(0)

        filename = f"{fan}.docx"
        return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=filename)
    except Exception as e:
            return render_template('error.html', error_message=str(e))  # Handle errors

@app.route('/generate_and_download_ru', methods=['POST'])
def generate_and_download_ru():
    if request.method == 'POST':
        kafedra_nomi = request.form['kafedra_nomi']
        guruh = request.form['guruh']
        fan = request.form['fan']
        semestr = request.form['semestr']
        tuzuvchi = request.form['tuzuvchi']
        zav_kaf = request.form['zav_kaf']
        kerakli_variantlar_soni = int(request.form['kerakli_variantlar_soni'])
        variantda_test_soni = int(request.form['variantda_test_soni'])
        fayl = request.files['savol_file']
        
        try:
            zip_data = generate_docx_files_ru(
                kafedra_nomi=kafedra_nomi,
                guruh =guruh,
                fan =fan,
                semestr =semestr,
                tuzuvchi =tuzuvchi,
                zav_kaf =zav_kaf,
                kerakli_variantlar_soni =kerakli_variantlar_soni,
                variantda_test_soni =variantda_test_soni,
                fayl = fayl
            )
            response = make_response(zip_data)
            response.headers['Content-Type'] = 'application/zip'
            response.headers['Content-Disposition'] = f'attachment; filename={fan}_test.zip'
            return response
        except Exception as e:
            return jsonify({'error': 'Failed to generate DOCX files', 'Error generating DOCX files':f'{e}'}), 500

@app.route('/generate_docx', methods=['POST'])
def generate_docx():
    try:
        # Get form data
        kafedra_nomi = request.form['kafedra_nomi']
        fan = request.form['fanimiz']
        semestr = request.form['semestr']
        tuzuvchi = request.form['tuzuvchi']
        kafedra_mudiri = request.form['kafedra_mudiri']
        bilet_soni = int(request.form['bilet_soni'])
        savollar_soni = int(request.form['savollar_soni'])
        # Get uploaded Excel file
        savol_file = request.files['savol_file']
        if savol_file.filename != '':
            savollar = savollar_from_xls(savol_file)
        
        document = Document()
        style = document.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)


        for i in range(bilet_soni):
            p_v1 = document.add_paragraph()
            p_v2 = document.add_paragraph()
            p_v3 = document.add_paragraph()
            
            v1 = p_v1.add_run('Namangan muhandislik - qurilish instituti\n')
            v2 = p_v1.add_run(f'«{kafedra_nomi}» kafedrasi \n Oraliq nazorat uchun savollar\n')
            v2 = p_v1.add_run(f'«{fan}» fanidan ({semestr}-semestr uchun)\n')
            v3 = p_v1.add_run(f'Oraliq nazorat savollari\n')
            v4 = p_v1.add_run(f'{i+1} - variant')
        
            for i in range(savollar_soni):
                p_v2.add_run(f"{i+1}) {random.choice(savollar)}\n")

            
            p_v1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_v2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_v3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            v1.bold = True
            v2.bold = True
            v3.bold = True
            v4.bold = True
            table = document.add_table(rows=2, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cell1 = table.rows[0].cells
            
            hdr_cell1[0].text = 'Tuzuvchi:'
            hdr_cell1[1].text = ''
            hdr_cell1[2].text = f'{tuzuvchi}'
            hdr_cell2 = table.rows[1].cells
            hdr_cell2[0].text = 'Kafedra mudiri:'
            hdr_cell2[1].text = ''
            hdr_cell2[2].text = f'{kafedra_mudiri}'
            # p_v3.add_run("\n")
        
        
        output = BytesIO()
        document.save(output)
        output.seek(0)

        filename = f"{fan}.docx"
        return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=filename)
    except Exception as e:
            return render_template('error.html', error_message=str(e))  # Handle errors

@app.route('/generate_and_download', methods=['POST'])
def generate_and_download():
    if request.method == 'POST':
        kafedra_nomi = request.form['kafedra_nomi']
        guruh = request.form['guruh']
        fan = request.form['fan']
        semestr = request.form['semestr']
        tuzuvchi = request.form['tuzuvchi']
        zav_kaf = request.form['zav_kaf']
        kerakli_variantlar_soni = int(request.form['kerakli_variantlar_soni'])
        variantda_test_soni = int(request.form['variantda_test_soni'])
        fayl = request.files['savol_file']
        
        try:
            zip_data = generate_docx_files(
                kafedra_nomi=kafedra_nomi,
                guruh =guruh,
                fan =fan,
                semestr =semestr,
                tuzuvchi =tuzuvchi,
                zav_kaf =zav_kaf,
                kerakli_variantlar_soni =kerakli_variantlar_soni,
                variantda_test_soni =variantda_test_soni,
                fayl = fayl
            )
            response = make_response(zip_data)
            response.headers['Content-Type'] = 'application/zip'
            response.headers['Content-Disposition'] = f'attachment; filename={fan}_test.zip'
            return response
        except Exception as e:
            return jsonify({'error': 'Failed to generate DOCX files', 'Error generating DOCX files':f'{e}'}), 500

 
if __name__ == '__main__':
    app.run(debug=True, port=8080)
